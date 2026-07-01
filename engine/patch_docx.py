#!/usr/bin/env python3
"""Patch review tokens in a generated DOCX with AI-provided values.

Usage called from JS/Pyodide:
    patch_docx(docx_path, overrides_json_path, output_path)

Where overrides JSON is: {"Field Name": "AI-found value"}
"""
import os, sys, json
import docx

def patch_review_tokens(docx_path, overrides, output_path):
    d = docx.Document(docx_path)
    patched = 0
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        text = run.text
                        for field, value in overrides.items():
                            # Match any ‹REVIEW — ...› token
                            if '‹REVIEW' in text:
                                # Replace the entire review token with the AI value
                                import re
                                new_text = re.sub(r'‹REVIEW[^›]*›', value, text)
                                if new_text != text:
                                    run.text = new_text
                                    text = new_text
                                    patched += 1
                                    break
    d.save(output_path)
    return patched

if __name__ == '__main__':
    docx_path = sys.argv[1]
    overrides_path = sys.argv[2]
    output_path = sys.argv[3]
    with open(overrides_path) as f:
        overrides = json.load(f)
    n = patch_review_tokens(docx_path, overrides, output_path)
    print(f'Patched {n} review tokens')
    with open(output_path.replace('.docx', '_ai_report.json'), 'w') as f:
        json.dump({'patched': n, 'overrides': overrides}, f)
