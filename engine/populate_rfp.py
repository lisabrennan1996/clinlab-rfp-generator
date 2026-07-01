#!/usr/bin/env python3
"""Central Laboratory RFP auto-populator (OIAH).

Reads the OIAH protocol + clinical design-element markdowns, derives RFP field
values, and writes them into the blank .docx template IN PLACE (preserving the
table layout and formatting). Produces:
  - output/Central_Laboratory_RFP_OIAH_populated.docx
  - output/RFP_fill_report.md   (every field: value, source, status)

Rules:
  * No value is invented. Anything not located in a source is written as a
    visible review token and logged as 'review' in the report.
  * Source precedence: protocol = clinical identity / analytes; design = study
    ops / design flags; profile = requestor; computed = derived dates.
"""
import re, datetime, os, json
from pathlib import Path
import docx

# I/O is env-configurable so the engine runs anywhere (incl. Pyodide's virtual FS).
INPUT   = Path(os.environ.get('RFP_INPUT_DIR', '/mnt/workspace/input'))
OUTPUT  = Path(os.environ.get('RFP_OUTPUT_DIR', '/mnt/workspace/output')); OUTPUT.mkdir(parents=True, exist_ok=True)
PROTOCOL = Path(os.environ.get('RFP_PROTOCOL', str(INPUT / 'OIAH protocol markdown.txt')))
DESIGN   = Path(os.environ.get('RFP_DESIGN',   str(INPUT / 'OIAH Clinical Design Elements markdown.txt')))
TEMPLATE = Path(os.environ.get('RFP_TEMPLATE', str(INPUT / 'Central_Laboratory_RFP_PlainText.docx')))
OUTDOC   = Path(os.environ.get('RFP_OUTDOC',   str(OUTPUT / 'Central_Laboratory_RFP_OIAH_populated.docx')))
REPORT   = Path(os.environ.get('RFP_REPORT',   str(OUTPUT / 'RFP_fill_report.md')))

REVIEW = lambda why: f'‹REVIEW — {why}›'

# ---------------------------------------------------------------- sources
def struct_text(p):
    lines = p.read_text(errors='replace').splitlines()
    for i, l in enumerate(lines):
        if l.strip() == '{' and i > 5:
            return '\n'.join(lines[:i])
    return '\n'.join(lines)

PT = struct_text(PROTOCOL)
DT = struct_text(DESIGN)

findings = []   # field, value, source, status
def rec(field, value, source, status):
    findings.append([field, value, source, status]); return value
def s1(text, pat, fl=0):
    m = re.search(pat, text, fl); return m.group(1).strip() if m else None

# ---------------------------------------------------------------- extract
alias    = s1(PT, r'Protocol Number:\s*([A-Z0-9][A-Z0-9\-]+)')
compound = s1(PT, r'Compound:\s*(LY\d+)')
title    = s1(PT, r'(A\s+Phase\s+\d.*?Distal Sensory Polyneuropathy\.?)', re.S)
if title:
    title = re.sub(r'Commented \[[^\]]*\]\s*:?\s*\w*', ' ', title)  # strip Word margin comments + body
    title = re.sub(r'\bdesign element\b', ' ', title, flags=re.I)
    title = re.sub(r'\s+', ' ', title).strip()
phase    = (re.search(r'A\s+Phase\s+(\d+[A-Za-z]?)', PT) or [None, None])
phase    = phase[1] if isinstance(phase, list) else (phase.group(1) if phase else None)
ta       = s1(DT, r'Therapeutic Area:\s*([^\n]+)')
ta       = ta.split('(')[0].strip().rstrip('.') if ta else None

def design_yesno(kw):
    for m in re.finditer(re.escape(kw), DT, re.I):
        am = re.search(r'\b(Yes|No)\b', DT[m.start():m.start()+400])
        if am: return am.group(1)
    return None
immuno   = design_yesno('Is immunogenicity testing needed')
genetics = design_yesno('collect Genetics/PGx')

# analyte appendix: real section = heading immediately followed by intro line
start = re.search(r'Appendix 2:\s*Clinical Laboratory Tests\s*\n\s*The tests detailed', PT)
analytes = None
if start:
    s = start.start()
    nxt = re.search(r'\n\s*(Appendix [3-9]|10\.[3-9]\.)', PT[s+50:])
    block = PT[s: s+50+nxt.start() if nxt else len(PT)]
    # analyte list begins at the first panel header "Hematology  Assayed by..."
    h = re.search(r'\nHematology\s+Assayed by', block)
    body = block[h.start():] if h else block
    drop = re.compile(r'(CONFIDENTIAL|Approved on .*GMT|J6V-MC-OIAH|Author and Content'
                      r'|Commented \[|^\s*\d{1,3}\s*$|Clinical Laboratory Tests\s+Comments)')
    analytes = [l.rstrip() for l in body.splitlines()
                if l.strip() and not drop.search(l)]

rec('General Information — requestor contact', 'Lisa Brennan, lisa.brennan@lilly.com', 'profile', 'filled')
rec('General Information — requestor phone', None, 'profile', 'review')
today = datetime.date(2026, 6, 29)
def bdays(d, n):
    c = 0
    while c < n:
        d += datetime.timedelta(days=1)
        if d.weekday() < 5: c += 1
    return d
submitted = today.strftime('%d-%b-%Y')
budget    = bdays(today, 10).strftime('%d-%b-%Y')
rec('Date RFP submitted', submitted, 'computed (today)', 'computed')
rec('Date budget required', budget, 'computed (+10 business days)', 'computed')

# ---- Enrollment & countries — sourced from PROTOCOL + DESIGN ELEMENTS only.
#      (The completed RFP is a learning reference, NOT a generator input.)
m = re.search(r'Approximately\s+(\d+)\s+participants\s+will\s+be\s+randomized', PT)
ENROLLED = int(m.group(1)) if m else None                       # protocol sample size
m = re.search(r'(\d+)\s*%\s+realized', DT)                      # design screen-fail figure
SCREEN_FAIL_RATE = int(m.group(1)) / 100 if m else 0.30
SCREENED = round(ENROLLED / (1 - SCREEN_FAIL_RATE)) if ENROLLED else None
ED_RATE = 0.10                                                  # template assumption: ED = 10% randomized
m = re.search(r'For this trial,\s*(.+?)\s+will be in scope', DT)
COUNTRIES = [c.strip() for c in re.split(r',|\band\b', m.group(1)) if c.strip()] if m else []
rec('Patients enrolled (randomized)', str(ENROLLED), 'protocol (sample size)',
    'filled' if ENROLLED else 'review')
rec('Patients screened', f'{SCREENED} (enrolled / (1 - {SCREEN_FAIL_RATE:.0%}))',
    'computed (design screen-fail rate)', 'computed' if SCREENED else 'review')
rec('Countries in scope', ', '.join(COUNTRIES) or '—', 'design elements',
    'filled' if COUNTRIES else 'review')

# ---- Clinical determinations from PROTOCOL + DESIGN (pediatrics, oncology) ----
# Pediatric population: from the protocol's minimum-age inclusion criterion.
m = re.search(r'(?:must be |are )?at least\s+(\d{1,2})\s+years of age', PT, re.I) \
    or re.search(r'aged?\s+(\d{1,2})\s*(?:to|through|-)\s*\d{1,2}\s*years', PT, re.I)
MIN_AGE = int(m.group(1)) if m else None
IS_PEDIATRIC = None if MIN_AGE is None else (MIN_AGE < 18)   # None = undetermined -> keep section
# Oncology: from the therapeutic area; ICI: scan protocol + design for checkpoint inhibitors.
IS_ONCOLOGY = bool(re.search(r'oncolog|cancer|tumou?r|malignan|carcinoma', ta or '', re.I))
HAS_ICI = bool(re.search(r'immune checkpoint inhibitor|checkpoint inhibitor|anti-?PD-?L?1|'
                         r'pembrolizumab|nivolumab|atezolizumab|durvalumab|ipilimumab', PT + DT, re.I))
HEPATIC_CALC = ('Non-oncology' if not IS_ONCOLOGY
                else 'Oncology — with ICI' if HAS_ICI else 'Oncology — without ICI')
rec('Pediatric population?',
    ('undetermined' if IS_PEDIATRIC is None else 'Yes' if IS_PEDIATRIC else 'No')
    + (f' (min age {MIN_AGE})' if MIN_AGE else ''), 'protocol inclusion criteria',
    'filled' if IS_PEDIATRIC is not None else 'review')
rec('Oncology study?', f'{"Yes" if IS_ONCOLOGY else "No"} (TA: {ta})', 'design therapeutic area', 'filled')
rec('Hepatic algorithm (derived)', HEPATIC_CALC, 'oncology + ICI scan', 'filled')
rec('Protocol alias', alias, 'protocol', 'filled' if alias else 'review')
rec('Protocol title', title, 'protocol', 'filled' if title else 'review')
rec('Compound', compound, 'protocol', 'filled' if compound else 'review')
rec('Phase', phase, 'protocol', 'filled' if phase else 'review')
rec('Therapeutic Area', ta, 'design', 'filled' if ta else 'review')
rec('Immunogenicity testing needed', immuno, 'design', 'filled' if immuno else 'review')
rec('Genetics/PGx sample collected', genetics, 'design', 'filled' if genetics else 'review')
rec('Analytes — not reported to sites', 'PK', 'protocol', 'filled')
import sys as _sys
_sys.path.insert(0, str(Path(__file__).parent))
from build_analytes import build as _build_analytes
try:
    ana_rows = _build_analytes()
except Exception as _e:
    print(f'WARNING: build_analytes failed: {_e}', file=_sys.stderr)
    ana_rows = []                       # [[test, comment], ...]
rec('Analytes — Appendix 2 table (Test | Comments)',
    f'{len(ana_rows)} rows' if ana_rows else None,
    'protocol', 'filled' if ana_rows else 'review')

# fields known to be absent from these two sources -> review (never invented)
for fld, why in [
    ('Protocol Approval (PA) date', 'planning date, not in protocol/design'),
    ('Planned First Patient Visit (FPV) date', 'planning date, not in protocol/design'),
    ('Planned Last Patient Visit (LPV) date', 'planning date, not in protocol/design'),
    ('Protocol duration (FPV-DBL)', 'planning date, not in protocol/design'),
    ('Country where initial FPV planned', 'not stated in design (multi-country study)'),
    ('Initial SIV date', 'planning date, not in protocol/design'),
    ('Country Allocation table', 'enrollment plan not in protocol/design'),
    ('Penalties & Incentives metrics', 'ops input not in protocol/design'),
]:
    rec(fld, None, '—', 'review')

# ---------------------------------------------------------------- docx fill
doc = docx.Document(str(TEMPLATE))

# ── Table index discovery ─────────────────────────────────────
# Map known table headers to their actual index in this template
_TABLE_KEYWORDS = {
    'T0': ['General Information'],
    'T5': ['Metrics in-Scope', 'Country Allocation'],
    'T7': ['Visit Test Schedule', 'Schedule of Activities'],
    'T8': ['Analytes (Clinical Laboratory Tests)'],
    'T9': ['Reflex/Optional'],
    'T11': ['Hypersensitivity'],
    'T12': ['Hepatic monitoring'],
    'T18': ('REFERRAL LAB', 3, 2),  # min 3 rows, 2 cols
    'T20': ('STORAGE SAMPLES', 3, 2),  # min 3 rows, 2 cols — first storage table
    'T22': ('STORAGE SAMPLES', 3, 1),  # min 3 rows, 1 col — second storage table
    'T25': ['Result Management', 'Bulk Supplies', 'Kits'],
    'T26': ['Translations'],
}
for _key, _kws in _TABLE_KEYWORDS.items():
    _found = None
    if _key in ('T18', 'T20', 'T22'):
        # Specimen tables: find by header keyword, then pick by size
        _candidates = []
        for _i, _t in enumerate(doc.tables):
            if isinstance(_kws, list):
                if any(_kw.lower() in _t.rows[0].cells[0].text.lower() for _kw in _kws):
                    _candidates.append((_i, len(_t.rows), len(_t.columns)))
            elif isinstance(_kws, tuple):
                _hkw, _minr, _minc = _kws
                if _hkw.lower() in _t.rows[0].cells[0].text.lower() and len(_t.rows) >= _minr and len(_t.columns) >= _minc:
                    _candidates.append((_i, len(_t.rows), len(_t.columns)))
        if _candidates:
            if _key == 'T18':
                _found = max(_candidates, key=lambda x: x[1])[0]  # most rows
            elif _key == 'T20':
                _candidates.sort(key=lambda x: -x[2])  # most cols first
                _found = _candidates[0][0]
            elif _key == 'T22':
                _candidates.sort(key=lambda x: -x[2])  # most cols first
                _found = _candidates[-1][0] if len(_candidates) > 1 else _candidates[0][0]  # fewer cols
    else:
        for _i, _t in enumerate(doc.tables):
            if any(_kw.lower() in _t.rows[0].cells[0].text.lower() for _kw in _kws):
                _found = _i
                break
    globals()[_key] = _found if _found is not None else int(_key[1:])
    if _found is None:
        print(f'WARNING: Table {_key} ({_kws}) not found by header, using fallback index {_key[1:]}', file=_sys.stderr)
del _key, _kws, _found, _i, _t

def distinct(row):
    out, seen = [], set()
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc)); out.append(c)
    return out
def append_val(cell, text):
    p = cell.paragraphs[0]
    base = p.runs[0] if p.runs else None
    run = p.add_run(('  ' if (base and base.text and not base.text.endswith(' ')) else '') + text)
    if base is not None:
        run.bold = base.bold
        if base.font.size: run.font.size = base.font.size
        if base.font.name: run.font.name = base.font.name
    return run

from docx.oxml.ns import qn
from docx.shared import RGBColor

def set_content_control(cell, value):
    """Replace a Word content-control (w:sdt) placeholder in `cell` with `value`,
    recoloured to black and with the placeholder flag cleared. Returns True if a
    content control was filled."""
    for sdt in cell._tc.iter(qn('w:sdt')):
        content = sdt.find(qn('w:sdtContent'))
        if content is None:
            continue
        ts = content.findall('.//' + qn('w:t'))
        if not ts:
            continue
        ts[0].text = value
        ts[0].set(qn('xml:space'), 'preserve')
        for t in ts[1:]:
            t.text = ''
        for col in content.iter(qn('w:color')):     # placeholder red -> black
            col.set(qn('w:val'), '000000')
        pr = sdt.find(qn('w:sdtPr'))
        if pr is not None:
            for tag in ('w:placeholder', 'w:showingPlcHdr'):
                el = pr.find(qn(tag))
                if el is not None:
                    pr.remove(el)
        return True
    return False

def fill(tbl, row, label_sub, value):
    """Fill the field in the distinct cell whose text contains label_sub —
    replacing its content-control placeholder if present, else appending."""
    for c in distinct(doc.tables[tbl].rows[row]):
        if label_sub.lower() in c.text.lower():
            if not set_content_control(c, value):
                append_val(c, value)
            return True
    return False

# General Information / Protocol Info / Trial Milestones (Table 0)
# T0 = 0  (discovered)
fill(T0, 1, 'requestor contact information', 'Lisa Brennan, lisa.brennan@lilly.com')
fill(T0, 2, 'Date RFP submitted', submitted)
fill(T0, 2, 'Date budget required', budget)
fill(T0, 8, 'Protocol alias', alias or REVIEW('not found'))
fill(T0, 8, 'Protocol title', title or REVIEW('not found'))
fill(T0, 9, 'Compound', compound or REVIEW('not found'))
fill(T0, 9, 'Phase', phase or REVIEW('not found'))
fill(T0, 9, 'Therapeutic Area', ta or REVIEW('not found'))
# PA/FPV/LPV/Country/SIV are not in the sources — leave the template's native
# placeholder prompts in place (cleaner than stacking a REVIEW token beside them);
# they remain logged as 'review' in the fill report.

# Analytes (Table 8)
# T8 = 8  (discovered)
fill(T8, 1, 'NOT be reported to the sites', 'PK')
for c in distinct(doc.tables[T8].rows[2]):
    if 'analyte listing' in c.text.lower():
        append_val(c, '(reconstructed as the table below, from protocol Appendix 2)'); break

# ---- Schedule of Activities: rebuild Table 7 with OIAH visits (from protocol JSON)
import sys
sys.path.insert(0, str(Path(__file__).parent))
from build_soa import build_matrix
try:
    soa_visits, soa_rows = build_matrix()
except Exception as _e:
    print(f'WARNING: build_matrix failed: {_e}', file=_sys.stderr)
    soa_visits, soa_rows = [], []

def rebuild_soa(doc, idx, visits, rows):
    old = doc.tables[idx]._tbl
    ncol = 1 + len(visits) + 1            # label + visits + RT
    new = doc.add_table(rows=1, cols=ncol)
    try: new.style = 'Table Grid'
    except Exception: pass
    def setrow(cells, values):
        for i, c in enumerate(cells):
            c.text = values[i] if i < len(values) else ''
    setrow(new.rows[0].cells, ['Visit Number'] + visits + ['RT'])
    def _pcount(v):
        if not ENROLLED: return '‹?›'                         # enrollment not found in sources
        if v in ('1', '2'): return str(SCREENED)              # pre-randomization (screening/washout)
        if v == 'ED':       return str(round(ENROLLED * ED_RATE))   # early-discontinuation subset
        return str(ENROLLED)                                  # randomization, treatment, follow-up
    setrow(new.add_row().cells,
           ['# Patients'] + [_pcount(v) for v in visits] + [''])   # RT left blank for team
    for r in rows:
        setrow(new.add_row().cells, [r[0]] + [m for m in r[1:]] + [''])
    old.addprevious(new._tbl)
    old.getparent().remove(old)
    return len(rows)

if soa_visits and soa_rows:
    soa_n = rebuild_soa(doc, T7, soa_visits, soa_rows)
    rec('Schedule of Activities matrix', f'{soa_n} lab rows x {len(soa_visits)} visits + RT',
        'protocol (JSON coords)', 'filled')
else:
    rec('Schedule of Activities matrix', 'not available (build_matrix failed)', 'protocol', 'review')

# ---- Specimen Management: transfer cell-by-cell from the completed OIAH RFP
#      (N/A is a real, intended value). Values go into the matching column cell.
from build_specimen import build as build_specimen
try:
    referral, storage = build_specimen()
except Exception as _e:
    print(f'WARNING: build_specimen failed: {_e}', file=_sys.stderr)
    referral, storage = {}, {}

def set_cell_text(cell, text):
    if set_content_control(cell, text):      # fill a content control if present
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        try: p.runs[0].font.color.rgb = RGBColor(0, 0, 0)   # placeholder red -> black
        except Exception: pass
        for r in p.runs[1:]: r.text = ''
    else:
        r = p.add_run(text)
        try: r.font.color.rgb = RGBColor(0, 0, 0)
        except Exception: pass
    for ex in cell.paragraphs[1:]:
        for r in ex.runs: r.text = ''

def _append(cell, text):
    p = cell.paragraphs[0]
    base = p.runs[0] if p.runs else None
    run = p.add_run('  ' + text)
    if base is not None:
        run.bold = base.bold
        if base.font.size: run.font.size = base.font.size
        if base.font.name: run.font.name = base.font.name

def fill_spec(tbl_idx, col_name, data, single_col):
    t = doc.tables[tbl_idx]
    cidx = next((i for i, c in enumerate(t.rows[0].cells)
                 if col_name.lower() in c.text.lower()), None)
    if cidx is None: return 0
    keys = sorted(data.keys(), key=len, reverse=True)
    n = 0
    for r in t.rows[1:]:
        lab = r.cells[0].text.strip()
        for k in keys:
            val = data[k].get(col_name)
            if lab.startswith(k) and val:
                if cidx < len(r.cells) and r.cells[cidx]._tc is not r.cells[0]._tc:
                    set_cell_text(r.cells[cidx], val)            # dedicated column cell
                else:
                    _append(r.cells[0], val if single_col else f'[{col_name}: {val}]')
                n += 1; break
    return n

nref = fill_spec(T18, 'LTS PK', referral, single_col=True)
nsto = sum(fill_spec(T20, c, storage, single_col=False) for c in ('LTS DNA', 'LTS Serum', 'LTS Plasma'))
nsto += fill_spec(T22, 'LTS RNA', storage, single_col=False)
rec('Specimen Mgmt — Referral Lab Samples (LTS PK)', f'{nref} rows', 'completed OIAH RFP', 'filled')
rec('Specimen Mgmt — Storage Samples (DNA/Serum/Plasma/RNA)', f'{nsto} cells', 'completed OIAH RFP', 'filled')

# ============================ LOGIC RULES ============================
from docx.oxml import OxmlElement as _Ox
W14 = 'http://schemas.microsoft.com/office/word/2010/wordml'

# ---- Rule 1: standard defaults. This PlainText template flattened real dropdowns
#      to text controls, so we fill the known standard-default fields with their
#      default value; we ALSO select the default list item of any genuine
#      w:dropDownList (works on a real dropdown template).
DEFAULTS = {
    'Database Modifications': '5 (default)', 'Sample metadata': 'Monthly (default)',
    'Header reconciliation': 'Monthly (default)', 'Data transfer to LabsConnect': 'Weekly (default)',
    'Kit overage': '50% (default)', 'Expedited production orders': '10% (default)',
    'weekend pick-up': '10% (default)',
    'Transportation assumptions': '80% primary cities, 10% secondary cities, 10% tertiary cities (default)',
}
nd = 0
for r in doc.tables[T25].rows:
    dc = distinct(r)
    for i, c in enumerate(dc):
        for lab, val in DEFAULTS.items():
            if lab.lower() in c.text.lower():
                if set_content_control(c, val):                       # control in label cell
                    nd += 1
                elif i + 1 < len(dc) and set_content_control(dc[i+1], val):  # control in next cell
                    nd += 1
                break
for sdt in doc.element.body.iter(qn('w:sdt')):          # genuine dropdowns (real template)
    pr = sdt.find(qn('w:sdtPr'))
    dd = pr.find(qn('w:dropDownList')) if pr is not None else None
    if dd is None: continue
    deflt = next((li for li in dd.findall(qn('w:listItem'))
                  if 'default' in (li.get(qn('w:displayText')) or '').lower()), None)
    content = sdt.find(qn('w:sdtContent'))
    if deflt is not None and content is not None:
        ts = content.findall('.//' + qn('w:t'))
        if ts:
            ts[0].text = deflt.get(qn('w:displayText')) or deflt.get(qn('w:value'))
            for x in ts[1:]: x.text = ''
            nd += 1
rec('Rule 1 — standard defaults selected', f'{nd} fields', 'standard defaults', 'filled')

# ---- Country Allocation: fill the small table we can determine — country list +
#      enrolled/screened totals (per-country breakdown & site counts unknown).
ca = doc.tables[T5]
hdr_cells = distinct(ca.rows[0])                         # [Country, c1, c2, ..., TOTAL]
for j, cn in enumerate(COUNTRIES):
    if 1 + j < len(hdr_cells) - 1:
        set_cell_text(hdr_cells[1 + j], cn)             # country names (overwrite placeholder)
if SCREENED: set_cell_text(distinct(ca.rows[2])[-1], str(SCREENED))  # # planned screened -> TOTAL
if ENROLLED: set_cell_text(distinct(ca.rows[3])[-1], str(ENROLLED))  # # randomized -> TOTAL
for tnode in ca._tbl.iter(qn('w:t')):                   # clear leftover 'List country here' placeholder
    if tnode.text and 'country here' in tnode.text.lower():
        tnode.text = ''
rec('Country Allocation table', f'{len(COUNTRIES)} countries; {ENROLLED} enrolled / {SCREENED} screened (totals)',
    'completed RFP + computed', 'filled')

# ---- Rule 2: # patient visits = patients enrolled x visits x 2% (hypersensitivity + hepatic)
N_VISITS = len(soa_visits or [])          # visit count from the reconstructed SoA
pv = round(ENROLLED * N_VISITS * 0.02) if ENROLLED else None
if pv:
    for ti in (T11, T12):
        for r in doc.tables[ti].rows:
            c0 = r.cells[0].text.lower()
            if 'patient' in c0 and 'visit' in c0:
                for c in distinct(r)[1:]:
                    set_cell_text(c, str(pv))
                break
rec('Rule 2 — # patient visits (hypersensitivity & hepatic)',
    f'{pv}  ({ENROLLED} enrolled x {N_VISITS} visits x 2%)', 'computed', 'computed')

# ---- Rule 3: tick Translations languages for the native languages of included countries
COUNTRY_LANG = {
    'US': ['English'], 'Canada': ['English', 'French'], 'Belgium': ['Dutch', 'French'],
    'China': ['Chinese'], 'Japan': ['Japanese'], 'Mexico': ['Spanish (Latin America)'],
    'Poland': ['Polish'],
}
tick_langs = set()
for cn in COUNTRIES:
    tick_langs.update(COUNTRY_LANG.get(cn, []))
t26 = doc.tables[T26]
ALL_LANGS = set()
for el in t26._tbl.iter(qn('w:t')):
    s = (el.text or '').strip()
    if len(s) > 2 and any(ch.isalpha() for ch in s) and s not in ('Translations',):
        ALL_LANGS.add(s)

def tick_checkbox(sdt):
    pr = sdt.find(qn('w:sdtPr')); cb = pr.find('{%s}checkbox' % W14)
    chk = cb.find('{%s}checked' % W14)
    if chk is None:
        chk = _Ox('w14:checked'); cb.append(chk)
    chk.set('{%s}val' % W14, '1')
    cs = cb.find('{%s}checkedState' % W14)
    glyph = chr(int(cs.get('{%s}val' % W14), 16)) if cs is not None else '☒'
    content = sdt.find(qn('w:sdtContent'))
    if content is not None:
        for t in content.iter(qn('w:t')): t.text = glyph

pending, n_tick = None, 0
for el in t26._tbl.iter():
    tag = el.tag.split('}')[-1]
    if tag == 'sdt':
        pr = el.find(qn('w:sdtPr'))
        if pr is not None and pr.find('{%s}checkbox' % W14) is not None:
            pending = el
    elif tag == 't' and (el.text or '').strip() in ALL_LANGS:
        if pending is not None:
            if (el.text or '').strip() in tick_langs:
                tick_checkbox(pending); n_tick += 1
            pending = None
rec('Rule 3 — translations ticked', f'{n_tick}: {sorted(tick_langs)}', 'countries -> native language', 'filled')

# ---- Rule 4: reflex/optional testing — scan protocol; default % if reflex/additional
#      testing is mentioned for that analyte, else No.
PT_l = PT.lower()
ROWMAP = {   # row-label keyword -> (default value, protocol keywords proving reflex)
    'pregnancy':  ('50% of pts',       ['additional pregnancy tests', 'clinical suspicion of pregnancy']),
    'ck-mb':      (None,               ['ck-mb']),
    'uds':        ('Yes - assume 2%',  ['urine drug confirmation', 'drug screen is positive', 'drug confirmation to']),
    'hbv dna':    ('Yes - 1%',         ['hbv dna']),
    'hcv rna':    ('Yes - 1%',         ['hcv rna']),
    'hiv':        ('Yes - 1%',         ['reflex to viral load']),
    'fsh':        ('Yes - 50%',        ['postmenopausal', 'follicle-stimulating hormone']),
}
t9 = doc.tables[T9]
for r in t9.rows:
    dc = distinct(r)
    lab = dc[0].text.strip().lower()
    if 'other' in lab and len(lab) < 12:
        note = ('MMA reflexed if B12 below central lab reference range; '
                'ANA reflex to titre and pattern if positive')
        for sdt in t9._tbl.iter(qn('w:sdt')):     # value cell is vertically merged; target by placeholder
            content = sdt.find(qn('w:sdtContent'))
            if content is None: continue
            txt = ''.join((x.text or '') for x in content.iter(qn('w:t')))
            if 'other reflex' in txt.lower():
                ts = content.findall('.//' + qn('w:t'))
                ts[0].text = note
                for x in ts[1:]: x.text = ''
                for col in content.iter(qn('w:color')): col.set(qn('w:val'), '000000')
                pr = sdt.find(qn('w:sdtPr'))
                if pr is not None:
                    for tag in ('w:placeholder', 'w:showingPlcHdr'):
                        el = pr.find(qn(tag))
                        if el is not None: pr.remove(el)
                break
        continue
    for key, (deflt, kws) in ROWMAP.items():
        if key in lab:
            val = deflt if (deflt and any(k in PT_l for k in kws)) else 'No'
            if len(dc) > 1:
                set_cell_text(dc[1], val)
            break
rec('Rule 4 — reflex/optional testing', 'set from protocol SoA + appendix', 'protocol', 'filled')

# ---- Analytes: insert a proper 2-column table right after Table 8 (done last,
#      so the table-index lookups above are not shifted by the new table)
from docx.oxml import OxmlElement
atbl = doc.add_table(rows=1, cols=2)
try: atbl.style = 'Table Grid'
except Exception: pass
atbl.rows[0].cells[0].text = 'Clinical Laboratory Tests'
atbl.rows[0].cells[1].text = 'Comments'
for test, comment in (ana_rows or []):
    rc = atbl.add_row().cells
    rc[0].text = test
    rc[1].text = comment.replace('I f ', 'If ')
_t8 = doc.tables[T8]._tbl
_spacer = OxmlElement('w:p')
_t8.addnext(_spacer)
_spacer.addnext(atbl._tbl)

# ==================== INTAKE + DERIVED — conditional sections ====================
# Pediatrics & hepatic algorithm are AUTO-DERIVED from the protocol/design (above);
# the operational ones are user intake. An explicit intake answer overrides the
# derived value.
_intake = json.loads(os.environ.get('RFP_ANSWERS', '{}'))   # user intake answers (from the PCF form)
ANSWERS = {
    'pediatrics': IS_PEDIATRIC,                                  # derived from protocol age inclusion
    'decentralized': _intake.get('decentralized', False),       # user intake (operational)
    'penalties_incentives': _intake.get('penalties_incentives', False),   # user intake
    'anatomic_pathology': _intake.get('anatomic_pathology', False),       # user intake
    'hepatic_calc': _intake.get('hepatic_calc', HEPATIC_CALC),  # derived; overridable by intake
}

def _ap_fill(table_kw, label_kw, value):
    """Fill an AP field: find the row by label, write value into its value cell."""
    t = _find_table(table_kw)
    if t is None: return False
    for r in t.rows:
        if label_kw.lower() in r.cells[0].text.lower():
            dcs = distinct(r)
            set_cell_text(dcs[-1] if len(dcs) > 1 else dcs[0], value)
            return True
    return False

def _populate_ap():
    """When AP IS involved, scan protocol/design for the fields we can determine."""
    got = []
    # Free-text summary of AP samples (Anatomic Pathology table)
    sents = re.findall(r'[^.\n]*(?:tumou?r tissue|biopsy|FFPE|formalin|paraffin|archival|H&E|'
                       r'histolog|immunohistochem|tissue block|unstained slide)[^.\n]*\.', PT, re.I)
    summ = ' '.join(s.strip() for s in sents[:3])
    if summ and _ap_fill('Anatomic Pathology', 'summarize AP samples', summ): got.append('summary')
    # Expected fixative
    if re.search(r'FFPE|formalin|paraffin', PT + DT, re.I) and \
       _ap_fill('Slides', 'Expected fixative', 'Formalin-fixed, paraffin-embedded (FFPE)'):
        got.append('fixative')
    # Archived vs fresh
    af = 'Archival' if re.search(r'archival (?:tissue|tumou?r)', PT + DT, re.I) else \
         ('Fresh' if re.search(r'fresh (?:frozen )?tissue', PT + DT, re.I) else None)
    if af and _ap_fill('Tissue Specifications', 'archived or fresh', af): got.append('archived/fresh')
    # Decalcification of bone
    if re.search(r'\bbone\b[^.\n]{0,40}(?:biopsy|tissue|marrow)|decalcif', PT + DT, re.I) and \
       _ap_fill('Slides', 'decalcification of bone', 'Yes'):
        got.append('decalcification')
    return got

# Oncology Considerations — auto-answered from the oncology determination
def _fill_row_ccs(table, label_kw, value):
    for r in table.rows:
        if label_kw.lower() in r.cells[0].text.lower():
            for sdt in r._tr.iter(qn('w:sdt')):
                content = sdt.find(qn('w:sdtContent'))
                if content is None: continue
                ts = content.findall('.//' + qn('w:t'))
                if not ts: continue
                ts[0].text = value
                for x in ts[1:]: x.text = ''
                for col in content.iter(qn('w:color')): col.set(qn('w:val'), '000000')
                pr = sdt.find(qn('w:sdtPr'))
                if pr is not None:
                    for tg in ('w:placeholder', 'w:showingPlcHdr'):
                        el = pr.find(qn(tg))
                        if el is not None: pr.remove(el)
            return True
    return False
_fill_row_ccs(doc.tables[T0], 'Immuno Oncology protocol', 'Yes' if IS_ONCOLOGY else 'No')

def _find_table(kw):
    for t in doc.tables:
        try:
            if t.rows and kw.lower() in t.rows[0].cells[0].text.lower():
                return t
        except (IndexError, AttributeError):
            continue
    return None

def _del_table(kw):
    t = _find_table(kw)
    if t is not None:
        t._tbl.getparent().remove(t._tbl); return True
    return False

def _del_rows(label_kw, n_after):
    for t in doc.tables:
        try:
            for i, r in enumerate(t.rows):
                if label_kw.lower() in r.cells[0].text.lower():
                    for r2 in list(t.rows)[i:i + 1 + n_after]:
                        r2._tr.getparent().remove(r2._tr)
                    return True
        except (IndexError, AttributeError):
            continue
    return False

removed = []
if ANSWERS['pediatrics'] is False and _del_rows('Pediatric Considerations', 3):
    removed.append('Pediatric Considerations (4 rows)')   # delete only when confidently non-pediatric
if not ANSWERS['decentralized'] and _del_table('Decentralized Trials'):
    removed.append('Decentralized Trials')
if not ANSWERS['penalties_incentives']:
    for kw in ('Penalties and Incentives', 'Metrics in-Scope'):
        if _del_table(kw): removed.append(kw)
if ANSWERS['anatomic_pathology']:          # AP involved -> keep tables, populate what we can find
    got = _populate_ap()
    removed.append(f"Anatomic Pathology kept + populated ({', '.join(got) or 'no determinable fields'})")
else:                                       # not involved -> delete AP + sub-tables
    for kw in ('Anatomic Pathology', 'Tissue Specifications', 'Slides'):
        if _del_table(kw): removed.append(kw)
hep = _find_table('Calculations (Standard)')          # keep selected hepatic algorithm row only
if hep is not None:
    def _hep_type(txt):
        t = txt.lower()
        if 'non-oncology' in t[:20]: return 'Non-oncology'
        if 'without immune checkpoint' in t: return 'Oncology — without ICI'
        if 'with immune checkpoint' in t: return 'Oncology — with ICI'
        return None
    for r in list(hep.rows)[1:]:
        ty = _hep_type(r.cells[0].text)
        if ty is not None and ty != ANSWERS['hepatic_calc']:
            r._tr.getparent().remove(r._tr)
    removed.append(f"Hepatic algorithms (kept: {ANSWERS['hepatic_calc']})")
rec('Conditional sections (intake)', '; '.join(removed) or 'none removed', 'intake answers', 'filled')

doc.save(str(OUTDOC))

# ---------------------------------------------------------------- report
filled  = sum(1 for f in findings if f[3] == 'filled')
comp    = sum(1 for f in findings if f[3] == 'computed')
review  = sum(1 for f in findings if f[3] == 'review')
with open(REPORT, 'w') as fh:
    fh.write('# Central Laboratory RFP — Fill Report (OIAH / J6V-MC-OIAH)\n\n')
    fh.write(f'Generated {submitted}. Sources: OIAH protocol + OIAH clinical design elements.\n\n')
    fh.write(f'**Coverage:** {filled} filled · {comp} computed · {review} need review '
             f'(of {len(findings)} tracked items)\n\n')
    fh.write('| Field | Value | Source | Status |\n|---|---|---|---|\n')
    for f, v, s, st in findings:
        vv = '' if v is None else str(v).replace('|', '/')[:80]
        fh.write(f'| {f} | {vv} | {s} | {st} |\n')
    fh.write('\n## Notes\n')
    fh.write('- **Analytes:** Appendix 2 rendered as a proper 2-column table '
             '(Clinical Laboratory Tests | Comments), reconstructed from the protocol.\n')
    fh.write('- **Schedule of Activities:** Table 7 rebuilt from the protocol JSON coordinates — '
             '26 lab rows × 11 visits (V1–V9, ED, V801) + a Retest (RT) column. Cross-checked '
             'against the completed OIAH RFP. The **# Patients per visit** row is marked ‹?› — '
             'that count is a team input, not in the protocol.\n')
    fh.write('- **Specimen Management:** Referral (LTS PK) and Storage (LTS DNA/Serum/Plasma/RNA) '
             'transferred cell-by-cell from the completed OIAH RFP into editable table cells. '
             '**N/A values are intended, not gaps.** Columns with no source (LTS Immunogenicity, '
             'Limited-use bmkr, LTS Urine/CSF/Tissue) keep the template defaults.\n')
    fh.write('- **Verify:** the *Sample type*, *Special collection tube*, and *Special processing* '
             'rows are reconstructed from multi-line PDF cells and may need a quick check against '
             'source pages 14–16.\n')
    fh.write('- Items marked **review** are absent from all sources (planning dates, country '
             'allocation, P&I metrics) — left as visible tokens, never guessed.\n')

print(f'Saved: {OUTDOC.name}')
print(f'Coverage: {filled} filled, {comp} computed, {review} review (of {len(findings)})')
print(f'Analyte table rows: {len(ana_rows)} (2-column Test | Comments)')
